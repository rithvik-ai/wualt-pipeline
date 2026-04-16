"""Generate DISTRESS_ENGINE_DOCUMENTATION.docx from the markdown source."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re, os

doc = Document()

# ── Styles ─────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x10, 0x10, 0x20)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x44)

def bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# ════════════════════════════════════════════════════════════
# CONTENT
# ════════════════════════════════════════════════════════════

# Title page
doc.add_paragraph()
doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('WUALT')
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = RGBColor(0x6c, 0x63, 0xff)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run('Rule-Based Physiological\nDistress Detection Engine')
r2.font.size = Pt(22)
r2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run('Complete Technical Documentation')
r3.font.size = Pt(14)
r3.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = tp4.add_run('April 2026 · v1.0')
r4.font.size = Pt(11)
r4.font.color.rgb = RGBColor(0x88, 0x88, 0xa0)

doc.add_page_break()

# ─── 1. Executive Summary ─────────────────────────────────
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The WUALT Distress Detection Engine is a standalone, stateful, rule-based system '
    'that evaluates physiological signals from a wearable smart ring to detect stress '
    'and distress in real time.'
)
add_table(
    ['Property', 'Value'],
    [
        ['Architecture', 'Rule-based, deterministic, no ML'],
        ['Latency', '< 100 ms per evaluation (O(1) complexity)'],
        ['Input', 'Preprocessed pipeline output (4 signals + metadata)'],
        ['Output', 'State (normal/stress/distress) + confidence + alert'],
        ['Signals', 'Heart Rate, SpO2, HRV (RMSSD), Skin Temperature'],
        ['Statefulness', 'Persistence tracking, baseline warmup, alert rotation'],
        ['Dependencies', 'Python stdlib only — zero external packages'],
        ['File', 'distress_engine.py (991 lines, fully self-contained)'],
    ]
)
doc.add_paragraph(
    'The engine sits downstream of the preprocessing pipeline (preprocessing_pipeline.py), '
    'which handles raw sensor ingestion, EMA smoothing, Signal Quality Index (SQI) computation, '
    'clinical flag generation, and z-score baseline tracking. The distress engine consumes that '
    'processed output and makes a detection decision.'
)

# ─── 2. Why Rule-Based ────────────────────────────────────
doc.add_heading('2. Why Rule-Based, Not Machine Learning', level=1)
doc.add_paragraph('This is a deliberate architectural choice, not a limitation.')
add_table(
    ['Criterion', 'Rule-Based', 'ML-Based'],
    [
        ['Interpretability', 'Every decision traceable to a specific threshold', 'Black-box; SHAP/LIME are approximations'],
        ['Auditability', 'Clinicians can read and verify thresholds against literature', 'Requires statistical validation studies'],
        ['Regulatory path', 'Rules map directly to clinical guidelines (BTS/WHO)', 'Requires drift monitoring, retraining infrastructure'],
        ['Data requirements', 'Works immediately with zero training data', 'Requires large, labelled, diverse real-world datasets'],
        ['Failure modes', 'Predictable — change one number to fix', 'Unpredictable on unseen distributions'],
        ['Edge cases', 'Explicitly handled (exercise, COPD, pregnancy)', 'Must be represented in training data'],
        ['Latency', 'O(1) — constant time', 'Depends on model size'],
    ]
)
doc.add_paragraph(
    'For a safety-critical wearable device in early stage, rule-based detection provides '
    'a reliable, auditable, and immediately deployable baseline. ML can be layered on top '
    'later, but the rule-based engine remains as the safety net.'
)

# ─── 3. System Architecture ──────────────────────────────
doc.add_heading('3. System Architecture', level=1)
doc.add_paragraph(
    'The system has three layers: the smart ring hardware (sensors), the preprocessing pipeline '
    '(data cleaning, quality scoring, baseline tracking), and the distress detection engine '
    '(rule evaluation, state management, alert generation).'
)
bold_para('Critical separation:')
doc.add_paragraph(
    'The distress engine has ZERO imports from the preprocessing pipeline. It consumes a '
    'well-defined dictionary contract. This means either component can be replaced independently, '
    'the engine can be tested with synthetic data, and different hardware backends can feed '
    'the same engine.'
)

# ─── 4–5. I/O Contracts ──────────────────────────────────
doc.add_heading('4. Input Contract', level=1)
doc.add_paragraph(
    'The engine\'s evaluate() method accepts one argument — a dictionary with four top-level keys:'
)
add_table(
    ['Key', 'Type', 'Description'],
    [
        ['sample', 'dict', 'Processed sensor frame: HR, SpO2, HRV, temp, SQI, clinical_flags, accepted'],
        ['zscores', 'dict', 'Per-signal z-scores against personal baseline'],
        ['baseline_ready', 'bool', 'True after sufficient warmup frames (~30s)'],
        ['window', 'dict', 'Rolling 30-second window statistics (mean, var, min, max)'],
    ]
)
bold_para('Key fields the engine uses:')
doc.add_paragraph(
    'sample.hr, sample.spo2, sample.dyn_acc_mag, sample.temp (vital signs); '
    'sample.accepted (gate check); sample.sqi.overall (quality gate); '
    'sample.clinical_flags (advisory integration); '
    'zscores.hr/.hrv_rmssd/.spo2/.temp (personalized deviation); '
    'baseline_ready (cold-start handling); window.window_n (confidence scaling).'
)

doc.add_heading('5. Output Contract', level=1)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['state', '"normal" | "stress" | "distress"', 'Current detection state'],
        ['confidence', 'float 0.0–1.0', 'Engine confidence in the detection'],
        ['contributing_signals', 'list of str', 'Which signals triggered (e.g., ["hr", "spo2"])'],
        ['alert.title', 'str', 'Short heading (e.g., "Heart rate is up")'],
        ['alert.message', 'str', 'Full user-facing message — calm, no jargon'],
        ['alert.severity', '"low" | "medium" | "high"', 'Alert priority level'],
        ['debug.flags', 'dict of bool', 'Per-signal flag status'],
        ['debug.scores', 'dict of float', 'Per-signal severity scores (0–1)'],
        ['debug.persistence_s', 'float', 'Seconds this condition has persisted'],
        ['debug.motion_state', 'str', '"still" / "walking" / "active" / "exercise"'],
        ['debug.flag_count', 'int', 'Number of signals currently flagged'],
        ['debug.weighted_score', 'float', 'Combined weighted confidence score'],
        ['debug.source', 'str', '"zscore" / "absolute" / "clinical_flag"'],
    ]
)

# ─── 6. Seven-Stage Pipeline ─────────────────────────────
doc.add_heading('6. The Seven-Stage Evaluation Pipeline', level=1)
doc.add_paragraph('Each frame passes through seven sequential stages:')

doc.add_heading('6.1 Stage 1: Quality Gate', level=2)
doc.add_paragraph(
    'Two checks prevent decisions on bad data. (1) Frame acceptance: if sample.accepted == False '
    '(finger off, sensor failure, charging), return normal with confidence 0.0 and reset persistence. '
    '(2) SQI threshold: if sqi.overall < 0.5, data quality too low — same behavior.'
)
doc.add_paragraph(
    'Why reset persistence on gate failure: if a user removes their ring for 30 seconds during '
    'a stress episode, the persistence counter must not keep ticking. When they put it back on, '
    'the condition must re-establish from zero.'
)

doc.add_heading('6.2 Stage 2: Motion Classification', level=2)
doc.add_paragraph(
    'Determines whether the user is exercising so elevated HR can be suppressed as expected behavior.'
)
add_table(
    ['Motion State', 'dyn_acc_mag Range', 'Interpretation'],
    [
        ['still', '< 0.05 g', 'Sitting, lying, standing still'],
        ['walking', '0.05 – 0.25 g', 'Light walking, casual movement'],
        ['active', '0.25 – 0.50 g', 'Brisk walking, light sport'],
        ['exercise', '> 0.50 g', 'Running, HIIT, vigorous activity'],
    ]
)

doc.add_heading('6.3 Stage 3: Signal Flagging', level=2)
doc.add_paragraph(
    'The core detection step. For each of the four signal channels, three flag sources are '
    'checked in priority order: (1) Z-score thresholds (personalized, only when baseline_ready), '
    '(2) Absolute thresholds (universal, always active), (3) Clinical flags from preprocessing.'
)

bold_para('Z-Score Thresholds (Personalized):')
add_table(
    ['Signal', 'Threshold', 'Direction', 'Meaning'],
    [
        ['HR', 'z > +2.0', 'Elevated', 'HR is 2+ std devs above personal baseline'],
        ['HRV (RMSSD)', 'z < -1.5', 'Suppressed', 'HRV dropped 1.5+ std devs below baseline'],
        ['SpO2', 'z < -2.0', 'Suppressed', 'Oxygen 2+ std devs below baseline'],
        ['Temperature', 'z > +2.5', 'Elevated', 'Skin temp 2.5+ std devs above baseline'],
    ]
)

bold_para('Absolute Thresholds (Universal Safety Net):')
add_table(
    ['Signal', 'Threshold', 'Level', 'Clinical Basis'],
    [
        ['HR', '>= 120 bpm', 'High', 'Resting tachycardia'],
        ['HR', '>= 150 bpm', 'Emergency', 'Extreme; bypasses persistence'],
        ['SpO2', '<= 94%', 'Low', 'BTS/WHO clinical concern'],
        ['SpO2', '<= 90%', 'Emergency', 'Hypoxemia; bypasses persistence'],
        ['Skin Temp', '>= 37.8 C', 'High', 'Possible fever (skin is 4-7C below core)'],
    ]
)

bold_para('Clinical Flag Integration:')
add_table(
    ['Clinical Flag', 'Engine Action', 'Score Assigned'],
    [
        ['spo2_hypoxemia (XX%)', 'Flag SpO2', '0.85'],
        ['spo2_clinical_concern (XX%)', 'Flag SpO2', '0.55'],
        ['elevated_skin_temp (XXC)', 'Flag temp', '0.30'],
    ]
)

doc.add_heading('6.4 Stage 4: Motion-Aware Suppression', level=2)
doc.add_paragraph(
    'If motion is "active" or "exercise", AND only HR is flagged (not SpO2), '
    'AND HR < 150 bpm: the HR flag is completely cleared. Elevated HR during exercise is expected.'
)
bold_para('Critical exception:')
doc.add_paragraph(
    'If SpO2 is ALSO flagged during exercise, HR suppression does NOT occur. A runner with '
    'HR=145 and SpO2=91% is in potential danger — exercise does not explain the oxygen drop.'
)

doc.add_heading('6.5 Stage 5: Raw State Determination', level=2)
add_table(
    ['Flag Count', 'Raw State', 'Rationale'],
    [
        ['0 flags', 'normal', 'All signals within baseline'],
        ['1 flag (primary)', 'stress', 'Single signal deviation'],
        ['1 flag (temp only)', 'normal', 'Temp alone is too weak — advisory only'],
        ['2+ flags', 'distress', 'Multiple physiological systems disrupted'],
    ]
)

doc.add_heading('6.6 Stage 6: Persistence Enforcement', level=2)
doc.add_paragraph(
    'Conditions must persist for 60 seconds before the engine promotes the state. '
    'A 2-frame grace period absorbs sensor noise at the boundary.'
)
add_table(
    ['Configuration', 'Value', 'Effect'],
    [
        ['PERSISTENCE_STRESS_S', '60 seconds', 'Stress conditions must hold 60s before alerting'],
        ['PERSISTENCE_DISTRESS_S', '60 seconds', 'Distress conditions must hold 60s before escalating'],
        ['grace_frames', '2 frames', 'Brief signal dropouts do not reset the timer'],
    ]
)
bold_para('Emergency bypass:')
doc.add_paragraph(
    'SpO2 <= 90% or HR >= 150 at rest bypass persistence entirely and escalate immediately.'
)

doc.add_heading('6.7 Stage 7: Confidence Scoring', level=2)
bold_para('Normal state:')
doc.add_paragraph('confidence = SQI x 0.6 + window_factor x 0.4')
bold_para('Stress/Distress state:')
doc.add_paragraph('confidence = signal_score x 0.5 + persist_factor x 0.3 + sqi_factor x 0.2')
doc.add_paragraph(
    'Signal score is the weighted sum of flagged signals. Persist factor ramps from 0.3 to 1.0 '
    'over 120 seconds. SQI factor scales with data quality.'
)

# ─── 7. Three-State Model ────────────────────────────────
doc.add_heading('7. The Three-State Model', level=1)
add_table(
    ['State', 'Trigger', 'Severity', 'User Action'],
    [
        ['Normal', '0 primary signals flagged', 'Low', 'No intervention needed'],
        ['Stress', '1 primary signal sustained 60s', 'Low–Medium', 'Advisory: breathing, relaxation'],
        ['Distress', '2+ primary signals sustained 60s', 'High', 'Urgent: reach out for help'],
    ]
)
doc.add_paragraph(
    'Why 2+ signals for distress: a single elevated signal could have many benign explanations '
    '(coffee for HR, altitude for SpO2). When multiple systems are simultaneously disrupted, '
    'the probability of genuine distress is much higher.'
)

# ─── 8. Signal Channels ──────────────────────────────────
doc.add_heading('8. Signal Channels Deep Dive', level=1)

doc.add_heading('8.1 Heart Rate (HR) — Weight: 0.35', level=2)
doc.add_paragraph(
    'Primary stress indicator. Sympathetic activation directly elevates HR within seconds. '
    'Z-score threshold: z > +2.0. Absolute: >= 120 bpm (high), >= 150 bpm (emergency). '
    'Exercise suppression: YES — HR flag cleared during exercise unless SpO2 also drops.'
)

doc.add_heading('8.2 Blood Oxygen Saturation (SpO2) — Weight: 0.30', level=2)
doc.add_paragraph(
    'Critical safety signal. Normal: 95-100%. Below 94%: BTS/WHO clinical concern. '
    'Below 90%: hypoxemia / medical emergency (bypasses persistence). '
    'Exercise suppression: NEVER — SpO2 drop during exercise is always concerning.'
)

doc.add_heading('8.3 Heart Rate Variability (HRV RMSSD) — Weight: 0.25', level=2)
doc.add_paragraph(
    'Autonomic nervous system marker. Suppressed HRV indicates reduced parasympathetic tone. '
    'Z-score threshold: z < -1.5 (more sensitive than others because HRV captures chronic stress). '
    'No absolute threshold — HRV varies too much between individuals for universal limits. '
    'Exercise suppression: NOT suppressed — low HRV during exercise can indicate overtraining.'
)

doc.add_heading('8.4 Skin Temperature — Weight: 0.10 (weakest)', level=2)
doc.add_paragraph(
    'Weak supporting signal. Ring measures SKIN temperature (4-7C below core). '
    'Affected by ambient temperature, clothing, post-meal vasodilation. '
    'Z-score threshold: z > +2.5 (highest threshold). Absolute: >= 37.8C skin. '
    'SPECIAL RULE: Temperature ALONE never triggers stress — supporting signal only.'
)

# ─── 9. Dual-Threshold Architecture ──────────────────────
doc.add_heading('9. Dual-Threshold Architecture', level=1)
add_table(
    ['System', 'When Active', 'Advantage', 'Example'],
    [
        ['Z-Score (Personalized)', 'Only after baseline_ready', 'Catches deviations from individual normal', 'Athlete HR 48->80 bpm (z=+8) flagged even though 80 is "normal" for most'],
        ['Absolute (Universal)', 'Always, including cold-start', 'Safety net for extreme values', 'HR >= 150 or SpO2 <= 90% flagged regardless of baseline'],
    ]
)

# ─── 10. Clinical Flag Integration ───────────────────────
doc.add_heading('10. Clinical Flag Integration', level=1)
add_table(
    ['Concept', 'Reject Reasons', 'Clinical Flags'],
    [
        ['Purpose', 'Data quality failure', 'Health advisory'],
        ['Effect on frame', 'accepted=False, frame discarded', 'accepted unchanged, frame still analyzed'],
        ['Examples', 'finger_off, low_ppg_variance', 'spo2_clinical_concern, elevated_skin_temp'],
        ['Engine behavior', 'Frame skipped at gate', 'Flags integrated into signal flagging'],
    ]
)
doc.add_paragraph(
    'A person with fever and clean sensor data should still have their data analyzed. '
    'Clinical concerns are advisory — they never force frame rejection.'
)

# ─── 11. Exercise Suppression ────────────────────────────
doc.add_heading('11. Motion-Aware Exercise Suppression', level=1)
add_table(
    ['Scenario', 'HR', 'SpO2', 'Motion', 'HR Flag', 'Result'],
    [
        ['Normal jog', '140', '97%', 'exercise', 'CLEARED', 'Normal'],
        ['Sprint', '165', '97%', 'exercise', 'KEPT (>=150)', 'Stress'],
        ['Jog + altitude sickness', '140', '91%', 'exercise', 'KEPT (SpO2 also flagged)', 'Distress'],
        ['Resting stress', '120', '97%', 'still', 'KEPT', 'Stress'],
    ]
)

# ─── 12. Persistence Tracker ─────────────────────────────
doc.add_heading('12. Persistence Tracker', level=1)
doc.add_paragraph(
    'A single-frame anomaly should not generate an alert. The PersistenceTracker ensures '
    'conditions must be sustained for 60 seconds before state promotion.'
)
doc.add_paragraph(
    'The tracker maintains independent stress and distress timers with a 2-frame grace period. '
    'If a condition drops for 1-2 frames (sensor noise), the timer does NOT reset. Only after '
    '3+ consecutive clear frames does the timer reset.'
)
doc.add_paragraph(
    'The tracker fully resets on frame rejection (finger off, low SQI) to prevent phantom '
    'persistence from building up during data gaps.'
)

# ─── 13. Emergency Bypass ────────────────────────────────
doc.add_heading('13. Emergency Bypass Logic', level=1)
add_table(
    ['Condition', 'Action', 'Rationale'],
    [
        ['SpO2 <= 90%', 'Immediate DISTRESS (bypass persistence)', 'WHO hypoxemia definition — every second matters'],
        ['HR >= 150 at rest', 'Immediate STRESS minimum (bypass persistence)', 'Abnormal tachycardia while stationary'],
    ]
)
doc.add_paragraph(
    'The HR emergency bypass does NOT fire during exercise. HR 160 while sprinting is normal; '
    'HR 160 while sitting is not.'
)

# ─── 14. Confidence Scoring ──────────────────────────────
doc.add_heading('14. Confidence Scoring Model', level=1)
add_table(
    ['Scenario', 'Signal Score', 'Persistence', 'SQI', 'Confidence'],
    [
        ['HR only, just started', '0.21', '0.3', '1.0', '0.40'],
        ['HR + SpO2, 90 seconds', '0.49', '0.83', '1.0', '0.69'],
        ['HR + SpO2, 120+ seconds', '0.59', '1.0', '1.0', '0.80'],
        ['Emergency SpO2=88%', '0.45', '0.3', '1.0', '0.52'],
    ]
)

# ─── 15. Alert System ────────────────────────────────────
doc.add_heading('15. STAR-Principle Alert System', level=1)
add_table(
    ['Category', 'Severity', 'Example Message'],
    [
        ['NORMAL', 'Low', 'Everything looks steady. You\'re doing great.'],
        ['WARMING UP', 'Low', 'We\'re still learning your baseline.'],
        ['TEMP ELEVATED', 'Low', 'Your skin temperature is a bit warmer than usual.'],
        ['STRESS (HR)', 'Low', 'Your heart rate is a bit higher than usual.'],
        ['STRESS (SpO2)', 'Medium', 'Your oxygen level dipped a little.'],
        ['EXERCISE', 'Low', 'Your readings are elevated, but you seem to be moving.'],
        ['DISTRESS', 'High', 'We\'re noticing signs of distress.'],
        ['DISTRESS (SpO2)', 'High', 'Your oxygen level is lower than expected.'],
    ]
)
bold_para('Design principles:')
doc.add_paragraph(
    'No medical jargon ("oxygen level" not "SpO2 saturation"). No alarming language. '
    'Actionable suggestions (breathing exercises, reaching out). Calm, empathetic tone. '
    'No diagnosis — "signs of distress" not "panic attack". Multiple message variants '
    'rotate to avoid repetition.'
)

# ─── 16. Cold-Start ──────────────────────────────────────
doc.add_heading('16. Cold-Start Handling', level=1)
doc.add_paragraph(
    'When baseline_ready is False: z-score thresholds are disabled (all z-scores = 0.0), '
    'only absolute thresholds are active, the user sees "Getting to know you" message, '
    'and confidence is lower. Typical warmup: 30-60 seconds at 1 Hz.'
)
doc.add_paragraph(
    'Safety during cold-start: if someone puts on the ring and immediately has SpO2 at 88%, '
    'the absolute threshold catches it. The safety net is always on.'
)

# ─── 17. Synthetic Dataset ───────────────────────────────
doc.add_heading('17. Synthetic Training Dataset', level=1)
add_table(
    ['Metric', 'Value'],
    [
        ['Total frames', '9,926'],
        ['Scenario instances', '2,121'],
        ['User profiles', '15 (ages 16-65, M/F, sedentary to athletic, plus special populations)'],
        ['Scenario templates', '37 (7 normal, 4 exercise, 6 stress, 8 distress, 10 edge, 2 transition)'],
        ['Columns per frame', '53'],
        ['Normal : Stress : Distress : Rejected', '54.4% : 21.4% : 22.0% : 2.1%'],
    ]
)

# ─── 18. Validation Results ──────────────────────────────
doc.add_heading('18. Validation Results', level=1)
bold_para('Overall accuracy: 65.8%')
add_table(
    ['Class', 'Precision', 'Recall', 'F1'],
    [
        ['Normal', '94.6%', '60.9%', '74.1%'],
        ['Stress', '49.0%', '45.9%', '47.4%'],
        ['Distress', '48.3%', '93.9%', '63.8%'],
        ['Rejected', '100.0%', '100.0%', '100.0%'],
    ]
)
doc.add_paragraph(
    'Key takeaway: Distress recall is 93.9% — the engine catches almost all dangerous conditions. '
    'Normal precision is 94.6% — when it says normal, it\'s almost always right. '
    'The synthetic dataset reveals calibration gaps (exercise thresholds, HR-HRV coupling) '
    'that inform threshold tuning.'
)

# ─── 19. Calibration ─────────────────────────────────────
doc.add_heading('19. Calibration Levers', level=1)
add_table(
    ['Parameter', 'Current', 'Increase Effect', 'Decrease Effect'],
    [
        ['ZSCORE_TH["hr"]', '2.0', 'Fewer HR detections (more specific)', 'More HR detections (more sensitive)'],
        ['ZSCORE_TH["spo2"]', '-2.0', 'More SpO2 detections', 'Fewer SpO2 detections'],
        ['ABS_TH["hr_high"]', '120 bpm', 'Fewer absolute HR flags', 'More absolute HR flags'],
        ['ABS_TH["spo2_low"]', '94%', 'Fewer SpO2 flags', 'More SpO2 flags'],
        ['PERSISTENCE_STRESS_S', '60s', 'Slower to alert (fewer FPs)', 'Faster to alert'],
        ['MIN_SQI_FOR_DETECTION', '0.5', 'More frames skipped', 'More frames evaluated'],
        ['SIGNAL_WEIGHTS["hr"]', '0.35', 'HR has more confidence influence', 'HR has less influence'],
        ['MOTION_TH["active"]', '0.50', 'More classified as active', 'More classified as exercise'],
    ]
)

# ─── 20. File Reference ──────────────────────────────────
doc.add_heading('20. File Reference', level=1)
add_table(
    ['File', 'Purpose'],
    [
        ['distress_engine.py', 'Complete standalone distress detection engine (991 lines)'],
        ['generate_synthetic_dataset.py', 'Generates 9,926 labelled frames + runs validation'],
        ['synthetic_dataset.csv', 'Flat CSV with 53 columns — every frame labelled'],
        ['synthetic_scenarios.csv', 'Scenario-level summaries with mean vitals'],
        ['synthetic_dataset_summary.json', 'Full metadata, profiles, validation results'],
        ['preprocessing_pipeline.py', 'Upstream: raw sensor to processed output (separate system)'],
    ]
)

bold_para('How to run:')
add_code('python distress_engine.py                  # 13-scenario demo')
add_code('python generate_synthetic_dataset.py       # generate dataset + validate')
add_code('')
add_code('from distress_engine import DistressEngine')
add_code('engine = DistressEngine()')
add_code('result = engine.evaluate(pipeline_output)')
add_code('print(result["state"], result["alert"]["message"])')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of Document —')
r.font.color.rgb = RGBColor(0x88, 0x88, 0xa0)
r.font.size = Pt(10)

# ── Save ──────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'DISTRESS_ENGINE_DOCUMENTATION.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
